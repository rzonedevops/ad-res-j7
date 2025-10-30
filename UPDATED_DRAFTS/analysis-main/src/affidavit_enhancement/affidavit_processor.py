#!/usr/bin/env python3
"""
Affidavit Enhancement Processor

Automatically updates and refines affidavits based on the present state of the evidence base.
Monitors evidence files and applies relevant updates to legal documents while preserving
formatting and maintaining document integrity.
"""

import json
import logging
import re
import shutil
import time
import signal
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed, TimeoutError
import threading

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


class ProgressTracker:
    """Tracks progress of affidavit enhancement operations."""
    
    def __init__(self, total_items: int, operation_name: str = "Processing"):
        self.total_items = total_items
        self.completed_items = 0
        self.operation_name = operation_name
        self.start_time = time.time()
        self.lock = threading.Lock()
        self.current_item = ""
        self.last_log_time = 0
        
    def update(self, item_name: str = "", increment: int = 1):
        """Update progress with optional item name."""
        with self.lock:
            self.completed_items += increment
            if item_name:
                self.current_item = item_name
            self._log_progress()
            
    def _log_progress(self):
        """Log current progress with throttling to avoid spam."""
        current_time = time.time()
        
        # Throttle logging to avoid spam (max once per second)
        if current_time - self.last_log_time < 1.0 and self.completed_items < self.total_items:
            return
            
        self.last_log_time = current_time
        
        if self.total_items > 0:
            percentage = (self.completed_items / self.total_items) * 100
            elapsed = time.time() - self.start_time
            
            # Create a simple progress bar
            bar_length = 20
            filled_length = int(bar_length * self.completed_items // self.total_items)
            bar = '█' * filled_length + '░' * (bar_length - filled_length)
            
            if self.completed_items > 0:
                eta = (elapsed / self.completed_items) * (self.total_items - self.completed_items)
                eta_str = f"ETA: {eta:.1f}s"
            else:
                eta_str = "ETA: calculating..."
                
            status = f"{self.operation_name}: [{bar}] {self.completed_items}/{self.total_items} ({percentage:.1f}%) - {eta_str}"
            if self.current_item:
                status += f" - {self.current_item}"
                
            logger.info(status)
        else:
            logger.info(f"{self.operation_name}: {self.completed_items} items completed")


@dataclass
class AffidavitMetadata:
    """Metadata for affidavit documents."""
    
    file_path: str
    document_type: str  # "affidavit", "supporting_affidavit", "draft"
    case_number: Optional[str] = None
    deponent: Optional[str] = None
    last_enhanced: Optional[str] = None
    evidence_sources: List[str] = None
    enhancement_count: int = 0
    backup_path: Optional[str] = None
    
    def __post_init__(self):
        if self.evidence_sources is None:
            self.evidence_sources = []


@dataclass
class EvidenceUpdate:
    """Represents an evidence update that should be applied to affidavits."""
    
    evidence_file: str
    update_type: str  # "new_evidence", "correction", "enhancement"
    content: str
    priority: str  # "critical", "high", "medium", "low"
    applicable_sections: List[str]
    timestamp: str


@dataclass
class AffidavitSection:
    """Represents a section within an affidavit."""
    
    section_id: str
    title: str
    content: str
    paragraph_numbers: List[int]
    evidence_references: List[str]
    last_modified: str


@dataclass
class ReviewStatus:
    """Tracks review status for enhanced affidavits."""
    
    affidavit_path: str
    enhanced_path: str
    critical_updates: int
    high_priority_updates: int
    total_updates: int
    requires_review: bool
    review_completed: bool = False
    reviewer: Optional[str] = None
    review_date: Optional[str] = None
    approval_status: str = "pending"  # "pending", "approved", "rejected", "needs_revision"
    review_comments: str = ""
    legal_compliance_checked: bool = False
    ready_for_filing: bool = False


class AffidavitProcessor:
    """
    Processes and enhances affidavits based on evidence updates.
    """
    
    def __init__(self, config_path: Optional[str] = None):
        """
        Initialize the affidavit processor.
        
        Args:
            config_path: Path to configuration file
        """
        self.config = self._load_config(config_path)
        self.affidavit_dir = Path(self.config.get("affidavit_dir", "."))
        self.evidence_dir = Path(self.config.get("evidence_dir", "evidence"))
        self.backup_dir = Path(self.config.get("backup_dir", "backups"))
        self.output_dir = Path(self.config.get("output_dir", "enhanced_affidavits"))
        
        # Create directories
        for directory in [self.backup_dir, self.output_dir]:
            directory.mkdir(parents=True, exist_ok=True)
        
        self.affidavit_metadata: Dict[str, AffidavitMetadata] = {}
        self.evidence_updates: List[EvidenceUpdate] = []
        self.review_statuses: Dict[str, ReviewStatus] = {}
        
        # Review system configuration
        self.review_dir = Path(self.config.get("review_dir", "review_required"))
        self.review_dir.mkdir(parents=True, exist_ok=True)
        
    def _load_config(self, config_path: Optional[str]) -> Dict:
        """Load configuration from file."""
        default_config = {
            "affidavit_patterns": [
                "*AFFIDAVIT*.md",
                "*affidavit*.md", 
                "*AFFIDAVIT*.docx",
                "*affidavit*.docx"
            ],
            "evidence_patterns": [
                "evidence/**/*.md",
                "evidence/**/*.json",
                "case_*/analysis/*.md",
                "*EVIDENCE*.md"
            ],
            "critical_keywords": [
                "fraud", "perjury", "murder", "criminal", "evidence",
                "witness intimidation", "breach of fiduciary duty"
            ],
            "backup_on_change": True,
            "preserve_formatting": True,
            "auto_enhancement": True,
            "review_settings": {
                "require_review_for_critical": True,
                "require_review_for_multiple_high": True,
                "high_priority_threshold": 3,
                "auto_create_github_issue": True,
                "prevent_filing_without_review": True
            }
        }
        
        if config_path and Path(config_path).exists():
            with open(config_path, 'r') as f:
                user_config = json.load(f)
                default_config.update(user_config)
                
        return default_config
        
    def discover_affidavits(self) -> List[Path]:
        """
        Discover all affidavit files in the repository.
        
        Returns:
            List of affidavit file paths
        """
        affidavit_files = []
        
        for pattern in self.config["affidavit_patterns"]:
            affidavit_files.extend(self.affidavit_dir.rglob(pattern))
            
        logger.info(f"Discovered {len(affidavit_files)} affidavit files")
        return affidavit_files
        
    def analyze_evidence_changes(self, since_timestamp: Optional[str] = None) -> List[EvidenceUpdate]:
        """
        Analyze evidence directory for changes that should update affidavits.
        
        Args:
            since_timestamp: Only consider changes after this timestamp
            
        Returns:
            List of evidence updates
        """
        updates = []
        
        for pattern in self.config["evidence_patterns"]:
            for evidence_file in self.affidavit_dir.rglob(pattern.split('/')[-1]):
                if self._is_relevant_evidence(evidence_file):
                    update = self._extract_evidence_update(evidence_file, since_timestamp)
                    if update:
                        updates.append(update)
                        
        logger.info(f"Found {len(updates)} evidence updates")
        return updates
        
    def _is_relevant_evidence(self, evidence_file: Path) -> bool:
        """
        Determine if an evidence file is relevant for affidavit updates.
        
        Args:
            evidence_file: Path to evidence file
            
        Returns:
            True if relevant
        """
        try:
            content = evidence_file.read_text(encoding='utf-8')
            
            # Check for critical keywords
            for keyword in self.config["critical_keywords"]:
                if keyword.lower() in content.lower():
                    return True
                    
            # Check for specific patterns
            critical_patterns = [
                r'criminal\s+proceedings?',
                r'evidence\s+of\s+fraud',
                r'new\s+evidence',
                r'witness\s+statement',
                r'expert\s+opinion'
            ]
            
            for pattern in critical_patterns:
                if re.search(pattern, content, re.IGNORECASE):
                    return True
                    
        except Exception as e:
            logger.warning(f"Error reading evidence file {evidence_file}: {e}")
            
        return False
        
    def _extract_evidence_update(self, evidence_file: Path, since_timestamp: Optional[str]) -> Optional[EvidenceUpdate]:
        """
        Extract update information from an evidence file.
        
        Args:
            evidence_file: Path to evidence file
            since_timestamp: Only consider changes after this timestamp
            
        Returns:
            Evidence update or None
        """
        try:
            content = evidence_file.read_text(encoding='utf-8')
            modification_time = evidence_file.stat().st_mtime
            
            # Skip if file is older than since_timestamp
            if since_timestamp:
                since_dt = datetime.fromisoformat(since_timestamp.replace('Z', '+00:00'))
                file_dt = datetime.fromtimestamp(modification_time)
                if file_dt < since_dt:
                    return None
                    
            # Determine update priority based on content
            priority = self._determine_priority(content)
            
            # Extract applicable sections
            applicable_sections = self._extract_applicable_sections(content)
            
            # Determine update type
            update_type = self._determine_update_type(content)
            
            return EvidenceUpdate(
                evidence_file=str(evidence_file),
                update_type=update_type,
                content=content,
                priority=priority,
                applicable_sections=applicable_sections,
                timestamp=datetime.fromtimestamp(modification_time).isoformat()
            )
            
        except Exception as e:
            logger.error(f"Error extracting evidence update from {evidence_file}: {e}")
            return None
            
    def _determine_priority(self, content: str) -> str:
        """Determine the priority of an evidence update."""
        critical_indicators = ['murder', 'criminal', 'perjury', 'urgent', 'critical']
        high_indicators = ['fraud', 'evidence', 'witness', 'breach']
        
        content_lower = content.lower()
        
        if any(indicator in content_lower for indicator in critical_indicators):
            return "critical"
        elif any(indicator in content_lower for indicator in high_indicators):
            return "high"
        else:
            return "medium"
            
    def _extract_applicable_sections(self, content: str) -> List[str]:
        """Extract which affidavit sections this evidence applies to."""
        sections = []
        
        # Look for section indicators
        section_patterns = {
            "background": r'background|introduction|history',
            "evidence": r'evidence|proof|documentation',
            "timeline": r'timeline|chronology|events',
            "financial": r'financial|money|funds|payment',
            "regulatory": r'regulatory|compliance|legal',
            "conclusion": r'conclusion|summary|prayer'
        }
        
        content_lower = content.lower()
        for section, pattern in section_patterns.items():
            if re.search(pattern, content_lower):
                sections.append(section)
                
        return sections
        
    def _determine_update_type(self, content: str) -> str:
        """Determine the type of update needed."""
        if re.search(r'new\s+evidence|additional\s+proo', content, re.IGNORECASE):
            return "new_evidence"
        elif re.search(r'correction|error|mistake|clarification', content, re.IGNORECASE):
            return "correction"
        else:
            return "enhancement"
            
    def enhance_affidavit(self, affidavit_path: Path, updates: List[EvidenceUpdate], timeout_seconds: int = 300) -> bool:
        """
        Enhance a specific affidavit with evidence updates.
        
        Args:
            affidavit_path: Path to affidavit file
            updates: List of evidence updates to apply
            timeout_seconds: Maximum time to spend on enhancement
            
        Returns:
            True if enhancement was successful
        """
        start_time = time.time()
        
        try:
            # Check file exists and is readable
            if not affidavit_path.exists():
                logger.error(f"Affidavit file not found: {affidavit_path}")
                return False
                
            if not affidavit_path.is_file():
                logger.error(f"Path is not a file: {affidavit_path}")
                return False
                
            # Check file size (warn if very large)
            file_size = affidavit_path.stat().st_size
            if file_size > 10 * 1024 * 1024:  # 10MB
                logger.warning(f"Large file detected ({file_size / 1024 / 1024:.1f}MB): {affidavit_path.name}")
                
            # Create backup
            if self.config.get("backup_on_change", True):
                backup_path = self._create_backup(affidavit_path)
                logger.debug(f"Created backup: {backup_path}")
                
            # Process based on file type
            if affidavit_path.suffix.lower() == '.docx':
                result = self._enhance_docx_affidavit(affidavit_path, updates)
            else:
                result = self._enhance_markdown_affidavit(affidavit_path, updates)
                
            # Log processing time
            processing_time = time.time() - start_time
            if processing_time > 30:  # Log if processing took more than 30 seconds
                logger.warning(f"Slow processing detected: {affidavit_path.name} took {processing_time:.1f}s")
            else:
                logger.debug(f"Processed {affidavit_path.name} in {processing_time:.1f}s")
                
            return result
                
        except FileNotFoundError as e:
            logger.error(f"File not found: {affidavit_path} - {e}")
            return False
        except PermissionError as e:
            logger.error(f"Permission denied: {affidavit_path} - {e}")
            return False
        except MemoryError as e:
            logger.error(f"Out of memory processing {affidavit_path}: {e}")
            return False
        except Exception as e:
            logger.error(f"Unexpected error enhancing affidavit {affidavit_path.name}: {e}")
            return False
            
    def _create_backup(self, file_path: Path) -> Path:
        """Create a backup of the original file."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_name = f"{file_path.stem}_backup_{timestamp}{file_path.suffix}"
        backup_path = self.backup_dir / backup_name
        
        shutil.copy2(file_path, backup_path)
        return backup_path
        
    def _enhance_markdown_affidavit(self, affidavit_path: Path, updates: List[EvidenceUpdate]) -> bool:
        """
        Enhance a markdown affidavit file.
        
        Args:
            affidavit_path: Path to markdown affidavit
            updates: List of evidence updates
            
        Returns:
            True if successful
        """
        try:
            # Read file with error handling
            try:
                content = affidavit_path.read_text(encoding='utf-8')
            except UnicodeDecodeError:
                # Try with different encoding
                content = affidavit_path.read_text(encoding='latin-1')
                logger.warning(f"Used latin-1 encoding for {affidavit_path.name}")
            except Exception as e:
                logger.error(f"Failed to read file {affidavit_path.name}: {e}")
                return False
                
            # Check content length
            if len(content) > 1_000_000:  # 1MB of text
                logger.warning(f"Very large content detected ({len(content)} chars): {affidavit_path.name}")
            
            # Sort updates by priority
            sorted_updates = sorted(updates, key=lambda x: {
                'critical': 0, 'high': 1, 'medium': 2, 'low': 3
            }.get(x.priority, 3))
            
            enhanced_content = content
            applied_updates = 0
            
            # Apply updates with progress tracking
            for i, update in enumerate(sorted_updates):
                try:
                    logger.debug(f"Applying update {i+1}/{len(sorted_updates)}: {update.update_type}")
                    enhanced_content = self._apply_markdown_update(enhanced_content, update)
                    applied_updates += 1
                except Exception as e:
                    logger.warning(f"Failed to apply update {i+1}: {e}")
                    continue
                
            # Add enhancement metadata
            try:
                enhanced_content = self._add_enhancement_metadata(enhanced_content, updates)
            except Exception as e:
                logger.warning(f"Failed to add metadata: {e}")
                # Continue without metadata rather than failing completely
                
            # Write enhanced version
            enhanced_path = self.output_dir / f"{affidavit_path.stem}_enhanced{affidavit_path.suffix}"
            
            try:
                enhanced_path.write_text(enhanced_content, encoding='utf-8')
            except Exception as e:
                logger.error(f"Failed to write enhanced file: {e}")
                return False
            
            logger.info(f"Enhanced markdown affidavit: {enhanced_path.name} ({applied_updates}/{len(updates)} updates applied)")
            return True
            
        except Exception as e:
            logger.error(f"Error enhancing markdown affidavit {affidavit_path.name}: {e}")
            return False
            
    def _enhance_docx_affidavit(self, affidavit_path: Path, updates: List[EvidenceUpdate]) -> bool:
        """
        Enhance a Word document affidavit.
        
        Args:
            affidavit_path: Path to Word document
            updates: List of evidence updates
            
        Returns:
            True if successful
        """
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available, skipping Word document enhancement")
            return False
            
        try:
            doc = Document(str(affidavit_path))
            
            # Sort updates by priority
            sorted_updates = sorted(updates, key=lambda x: {
                'critical': 0, 'high': 1, 'medium': 2, 'low': 3
            }.get(x.priority, 3))
            
            for update in sorted_updates:
                self._apply_docx_update(doc, update)
                
            # Save enhanced version
            enhanced_path = self.output_dir / f"{affidavit_path.stem}_enhanced{affidavit_path.suffix}"
            doc.save(str(enhanced_path))
            
            logger.info(f"Enhanced Word affidavit: {enhanced_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error enhancing Word affidavit: {e}")
            return False
            
    def _apply_markdown_update(self, content: str, update: EvidenceUpdate) -> str:
        """Apply an evidence update to markdown content."""
        
        # Find appropriate insertion point based on update type and sections
        if update.update_type == "new_evidence":
            return self._insert_new_evidence_markdown(content, update)
        elif update.update_type == "correction":
            return self._apply_correction_markdown(content, update)
        else:
            return self._enhance_existing_content_markdown(content, update)
            
    def _insert_new_evidence_markdown(self, content: str, update: EvidenceUpdate) -> str:
        """Insert new evidence into markdown affidavit."""
        
        # Extract key evidence points from the update
        evidence_summary = self._extract_evidence_summary(update.content)
        
        # Find insertion point (before conclusion or at end of evidence sections)
        insertion_patterns = [
            r'## CONCLUSION',
            r'## PRAYER',
            r'## WHEREFORE',
            r'---\s*$'
        ]
        
        insertion_point = len(content)
        for pattern in insertion_patterns:
            match = re.search(pattern, content, re.IGNORECASE | re.MULTILINE)
            if match:
                insertion_point = match.start()
                break
                
        # Create new evidence section
        timestamp = datetime.now().strftime("%Y-%m-%d")
        new_section = """
## ENHANCED EVIDENCE ANALYSIS (Added {timestamp})

### New Evidence from {Path(update.evidence_file).name}

{evidence_summary}

**Source**: {update.evidence_file}  
**Priority**: {update.priority.upper()}  
**Date Added**: {timestamp}

---

"""
        
        # Insert the new section
        enhanced_content = content[:insertion_point] + new_section + content[insertion_point:]
        return enhanced_content
        
    def _extract_evidence_summary(self, content: str) -> str:
        """Extract a concise summary of evidence from content."""
        
        # Look for key evidence indicators
        evidence_patterns = [
            r'evidence shows?.*?(?=\n\n|\n#|\n\*|$)',
            r'proof.*?(?=\n\n|\n#|\n\*|$)',
            r'documents? confirm.*?(?=\n\n|\n#|\n\*|$)',
            r'analysis reveals?.*?(?=\n\n|\n#|\n\*|$)'
        ]
        
        summaries = []
        for pattern in evidence_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE | re.DOTALL)
            summaries.extend(matches[:3])  # Limit to first 3 matches
            
        if summaries:
            return '\n'.join(f"- {summary.strip()}" for summary in summaries[:5])
        else:
            # Fallback to first few lines if no specific patterns found
            lines = content.split('\n')[:10]
            return '\n'.join(f"- {line.strip()}" for line in lines if line.strip())[:500] + "..."
            
    def _apply_correction_markdown(self, content: str, update: EvidenceUpdate) -> str:
        """Apply corrections to markdown content."""
        
        # Add correction notice
        timestamp = datetime.now().strftime("%Y-%m-%d")
        correction_notice = """
## CORRECTION NOTICE (Applied {timestamp})

**Source**: {Path(update.evidence_file).name}  
**Priority**: {update.priority.upper()}

**Correction Details**:
{self._extract_evidence_summary(update.content)}

---

"""
        
        # Find insertion point after introduction but before main evidence
        patterns = [
            r'## 1\.\s+INTRODUCTION.*?\n(?=## 2\.)',
            r'## INTRODUCTION.*?\n(?=## [A-Z])',
            r'do hereby make oath and state:.*?\n(?=---|\n## |\n\d+\.)'
        ]
        
        insertion_point = 0
        for pattern in patterns:
            match = re.search(pattern, content, re.IGNORECASE | re.DOTALL)
            if match:
                insertion_point = match.end()
                break
                
        enhanced_content = content[:insertion_point] + correction_notice + content[insertion_point:]
        return enhanced_content
        
    def _enhance_existing_content_markdown(self, content: str, update: EvidenceUpdate) -> str:
        """Enhance existing markdown content with additional information."""
        
        # Find relevant sections to enhance based on applicable_sections
        for section in update.applicable_sections:
            content = self._enhance_section_markdown(content, section, update)
            
        return content
        
    def _enhance_section_markdown(self, content: str, section_name: str, update: EvidenceUpdate) -> str:
        """Enhance a specific section with update information."""
        
        # Section patterns to find
        section_patterns = {
            'background': r'(## .*?BACKGROUND.*?\n)(.*?)(?=## |\n\n---|\Z)',
            'evidence': r'(## .*?EVIDENCE.*?\n)(.*?)(?=## |\n\n---|\Z)',
            'timeline': r'(## .*?TIMELINE.*?\n)(.*?)(?=## |\n\n---|\Z)',
            'financial': r'(## .*?FINANCIAL.*?\n)(.*?)(?=## |\n\n---|\Z)'
        }
        
        pattern = section_patterns.get(section_name.lower())
        if not pattern:
            return content
            
        match = re.search(pattern, content, re.IGNORECASE | re.DOTALL)
        if not match:
            return content
            
        section_header = match.group(1)
        section_content = match.group(2)
        
        # Add enhancement to section
        enhancement = """
### Additional Evidence ({datetime.now().strftime("%Y-%m-%d")})

{self._extract_evidence_summary(update.content)}

**Source**: {Path(update.evidence_file).name}  
"""
        
        enhanced_section = section_header + section_content.rstrip() + enhancement + "\n"
        
        # Replace in content
        enhanced_content = content.replace(match.group(0), enhanced_section)
        return enhanced_content
        
    def _apply_docx_update(self, doc, update: EvidenceUpdate):
        """Apply an evidence update to a Word document."""
        
        # Find insertion point
        if update.update_type == "new_evidence":
            self._insert_new_evidence_docx(doc, update)
        elif update.update_type == "correction":
            self._apply_correction_docx(doc, update)
        else:
            self._enhance_existing_content_docx(doc, update)
            
    def _insert_new_evidence_docx(self, doc, update: EvidenceUpdate):
        """Insert new evidence section in Word document."""
        
        # Find insertion point (before conclusion)
        insertion_index = len(doc.paragraphs)
        
        for i, paragraph in enumerate(doc.paragraphs):
            if any(keyword in paragraph.text.upper() for keyword in ['CONCLUSION', 'PRAYER', 'WHEREFORE']):
                insertion_index = i
                break
                
        # Add new evidence section
        timestamp = datetime.now().strftime("%Y-%m-%d")
        
        # Add section header
        header_para = doc.paragraphs[insertion_index]._element
        new_header = doc.add_paragraph()
        new_header.text = f"ENHANCED EVIDENCE ANALYSIS (Added {timestamp})"
        new_header.style = 'Heading 2'
        header_para.getparent().insert(insertion_index, new_header._element)
        
        # Add content
        content_para = doc.add_paragraph()
        content_para.text = f"New Evidence from {Path(update.evidence_file).name}\n\n"
        content_para.text += self._extract_evidence_summary(update.content)
        
    def _apply_correction_docx(self, doc, update: EvidenceUpdate):
        """Apply correction to Word document."""
        
        # Find insertion point after introduction
        insertion_index = 3  # Default after first few paragraphs
        
        for i, paragraph in enumerate(doc.paragraphs):
            if 'make oath and state' in paragraph.text.lower():
                insertion_index = i + 1
                break
                
        # Add correction notice
        timestamp = datetime.now().strftime("%Y-%m-%d")
        correction_para = doc.add_paragraph()
        correction_para.text = f"CORRECTION NOTICE (Applied {timestamp})\n\n"
        correction_para.text += f"Source: {Path(update.evidence_file).name}\n"
        correction_para.text += f"Priority: {update.priority.upper()}\n\n"
        correction_para.text += self._extract_evidence_summary(update.content)
        
    def _enhance_existing_content_docx(self, doc, update: EvidenceUpdate):
        """Enhance existing content in Word document."""
        
        # Find relevant paragraphs and add enhancements
        enhancement_text = f"\n\nAdditional Evidence ({datetime.now().strftime('%Y-%m-%d')}): "
        enhancement_text += self._extract_evidence_summary(update.content)[:200] + "..."
        
        # Add to end of document before conclusion
        for paragraph in reversed(doc.paragraphs):
            if not any(keyword in paragraph.text.upper() for keyword in ['CONCLUSION', 'PRAYER', 'WHEREFORE']):
                paragraph.text += enhancement_text
                break
                
    def _add_enhancement_metadata(self, content: str, updates: List[EvidenceUpdate]) -> str:
        """Add metadata about enhancements to the document."""
        
        metadata = """
---

## ENHANCEMENT METADATA

**Last Enhanced**: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}  
**Evidence Sources Processed**: {len(updates)}  
**Enhancement Count**: {len([u for u in updates if u.priority in ['critical', 'high']])} critical/high priority updates

### Evidence Sources:
"""
        
        for update in updates:
            metadata += f"- {Path(update.evidence_file).name} ({update.priority} priority, {update.update_type})\n"
            
        metadata += "\n---\n"
        
        return content + metadata
        
    def process_all_affidavits(self, since_timestamp: Optional[str] = None, max_workers: int = 2) -> Dict[str, bool]:
        """
        Process all discovered affidavits with available evidence updates.
        
        Args:
            since_timestamp: Only consider evidence changes after this timestamp
            max_workers: Maximum number of parallel workers for processing
            
        Returns:
            Dictionary of affidavit paths and their enhancement success status
        """
        results = {}
        
        logger.info("🔍 Starting affidavit enhancement process...")
        
        # Discover affidavits and evidence updates
        logger.info("📋 Discovering affidavit files...")
        affidavits = self.discover_affidavits()
        
        if not affidavits:
            logger.info("ℹ️ No affidavit files found")
            return results
            
        logger.info("📊 Analyzing evidence changes...")
        updates = self.analyze_evidence_changes(since_timestamp)
        
        if not updates:
            logger.info("ℹ️ No evidence updates found")
            return results
            
        logger.info(f"📝 Found {len(updates)} evidence updates to process")
        
        # Filter affidavits that have relevant updates
        affidavits_to_process = []
        for affidavit_path in affidavits:
            relevant_updates = self._filter_relevant_updates(affidavit_path, updates)
            if relevant_updates:
                affidavits_to_process.append((affidavit_path, relevant_updates))
            else:
                logger.info(f"⏭️ No relevant updates for {affidavit_path.name}")
                results[str(affidavit_path)] = True  # No updates needed, consider successful
                
        if not affidavits_to_process:
            logger.info("ℹ️ No affidavits need enhancement")
            return results
            
        # Initialize progress tracker
        progress = ProgressTracker(len(affidavits_to_process), "Enhancing Affidavits")
        
        # Process affidavits with progress tracking
        if max_workers > 1 and len(affidavits_to_process) > 1:
            # Use parallel processing for multiple affidavits
            results.update(self._process_affidavits_parallel(affidavits_to_process, progress, max_workers))
        else:
            # Use sequential processing
            results.update(self._process_affidavits_sequential(affidavits_to_process, progress))
            
        # Log final results
        successful = sum(1 for success in results.values() if success)
        total = len(results)
        logger.info(f"✅ Enhancement complete: {successful}/{total} affidavits processed successfully")
        
        return results
        
    def _process_with_timeout(self, func, timeout_seconds: int = 300, *args, **kwargs):
        """Execute a function with a timeout."""
        def timeout_handler(signum, frame):
            raise TimeoutError(f"Operation timed out after {timeout_seconds} seconds")
            
        # Set up timeout handler
        old_handler = signal.signal(signal.SIGALRM, timeout_handler)
        signal.alarm(timeout_seconds)
        
        try:
            result = func(*args, **kwargs)
            return result
        except TimeoutError:
            logger.error(f"Operation timed out after {timeout_seconds} seconds")
            return False
        finally:
            # Restore original handler
            signal.alarm(0)
            signal.signal(signal.SIGALRM, old_handler)
        
    def _process_affidavits_sequential(self, affidavits_to_process: List[Tuple[Path, List[EvidenceUpdate]]], 
                                     progress: ProgressTracker) -> Dict[str, bool]:
        """Process affidavits sequentially with progress tracking."""
        results = {}
        
        for affidavit_path, relevant_updates in affidavits_to_process:
            try:
                progress.update(f"Processing {affidavit_path.name}")
                
                # Use timeout wrapper for enhancement
                success = self._process_with_timeout(
                    self.enhance_affidavit, 
                    300,  # 5 minutes per affidavit
                    affidavit_path, 
                    relevant_updates
                )
                
                results[str(affidavit_path)] = success
                
                if success:
                    # Check if review is required for critical updates
                    enhanced_path = self.output_dir / f"{affidavit_path.stem}_enhanced{affidavit_path.suffix}"
                    review_status = self._create_review_status(affidavit_path, enhanced_path, relevant_updates)
                    
                    # Save review status if review is required
                    if review_status.requires_review:
                        self._save_review_status(review_status)
                        self.review_statuses[str(affidavit_path)] = review_status
                        
                        alert_msg = self._generate_review_alert(review_status)
                        logger.warning(alert_msg)
                        logger.warning(f"🚨 CRITICAL: {affidavit_path.name} requires legal review before filing!")
                    
                    logger.info(f"✅ Enhanced {affidavit_path.name}: Success")
                else:
                    logger.warning(f"⚠️ Enhanced {affidavit_path.name}: Failed")
                    
            except Exception as e:
                logger.error(f"❌ Error processing affidavit {affidavit_path.name}: {e}")
                results[str(affidavit_path)] = False
                
        return results
        
    def _process_affidavits_parallel(self, affidavits_to_process: List[Tuple[Path, List[EvidenceUpdate]]], 
                                   progress: ProgressTracker, max_workers: int) -> Dict[str, bool]:
        """Process affidavits in parallel with progress tracking."""
        results = {}
        
        def process_single_affidavit(affidavit_data):
            affidavit_path, relevant_updates = affidavit_data
            try:
                success = self.enhance_affidavit(affidavit_path, relevant_updates)
                
                if success:
                    # Check if review is required for critical updates
                    enhanced_path = self.output_dir / f"{affidavit_path.stem}_enhanced{affidavit_path.suffix}"
                    review_status = self._create_review_status(affidavit_path, enhanced_path, relevant_updates)
                    
                    # Save review status if review is required
                    if review_status.requires_review:
                        self._save_review_status(review_status)
                        self.review_statuses[str(affidavit_path)] = review_status
                        logger.warning(f"🚨 CRITICAL: {affidavit_path.name} requires legal review before filing!")
                
                progress.update(f"Completed {affidavit_path.name}")
                return str(affidavit_path), success, relevant_updates
            except Exception as e:
                logger.error(f"❌ Error processing affidavit {affidavit_path.name}: {e}")
                progress.update(f"Failed {affidavit_path.name}")
                return str(affidavit_path), False, []
                
        # Use ThreadPoolExecutor for parallel processing
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all tasks
            future_to_affidavit = {
                executor.submit(process_single_affidavit, affidavit_data): affidavit_data[0]
                for affidavit_data in affidavits_to_process
            }
            
            # Process completed tasks
            for future in as_completed(future_to_affidavit):
                try:
                    affidavit_path, success, relevant_updates = future.result()
                    results[affidavit_path] = success
                except Exception as e:
                    affidavit_path = future_to_affidavit[future]
                    logger.error(f"❌ Unexpected error processing {affidavit_path.name}: {e}")
                    results[str(affidavit_path)] = False
                    
        return results
        
    def _filter_relevant_updates(self, affidavit_path: Path, updates: List[EvidenceUpdate]) -> List[EvidenceUpdate]:
        """Filter updates that are relevant to a specific affidavit."""
        
        # For now, return all updates
        # In the future, this could be more sophisticated based on:
        # - Case numbers
        # - Parties involved  
        # - Subject matter
        # - Document relationships
        
        return [update for update in updates if update.priority in ['critical', 'high', 'medium']]
        
    def generate_enhancement_report(self, results: Dict[str, bool]) -> str:
        """
        Generate a report of enhancement activities.
        
        Args:
            results: Results from process_all_affidavits
            
        Returns:
            Enhancement report as string
        """
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        successful = sum(1 for success in results.values() if success)
        total = len(results)
        
        # Count documents requiring review
        reviews_required = sum(1 for status in self.review_statuses.values() if status.requires_review)
        total_critical = sum(status.critical_updates for status in self.review_statuses.values())
        total_high = sum(status.high_priority_updates for status in self.review_statuses.values())
        
        report = """# Affidavit Enhancement Report

**Generated**: {timestamp}  
**Total Affidavits Processed**: {total}  
**Successfully Enhanced**: {successful}  
**Failed**: {total - successful}

## Critical Review Alert

**Documents Requiring Legal Review**: {reviews_required}
**Total Critical Updates**: {total_critical}
**Total High Priority Updates**: {total_high}

## Enhancement Results

"""
        
        for affidavit_path, success in results.items():
            status = "✅ Success" if success else "❌ Failed"
            
            # Add review indicator
            review_indicator = ""
            if affidavit_path in self.review_statuses:
                review_status = self.review_statuses[affidavit_path]
                if review_status.requires_review:
                    review_indicator = " 🚨 **REQUIRES LEGAL REVIEW**"
            
            report += f"- `{Path(affidavit_path).name}`: {status}{review_indicator}\n"

        # Add review section if any documents require review
        if reviews_required > 0:
            report += """
## 🚨 DOCUMENTS REQUIRING LEGAL REVIEW

The following enhanced affidavits contain critical evidence updates and MUST be reviewed by qualified legal counsel before filing:

"""
            for affidavit_path, review_status in self.review_statuses.items():
                if review_status.requires_review:
                    report += """### {Path(affidavit_path).name}

- **Enhanced File**: `{Path(review_status.enhanced_path).name}`
- **Critical Updates**: {review_status.critical_updates}
- **High Priority Updates**: {review_status.high_priority_updates}
- **Total Updates**: {review_status.total_updates}
- **Status**: {review_status.approval_status.upper()}
- **Ready for Filing**: {'✅ YES' if review_status.ready_for_filing else '❌ NO'}

"""
            
        report += """
## Enhanced Files Location

Enhanced affidavits have been saved to: `{self.output_dir}`

## Backup Files Location

Original files backed up to: `{self.backup_dir}`

## Review Status Files

Review tracking files are saved in: `{self.review_dir}`

## Next Steps for Documents Requiring Review

1. Review enhanced affidavits for accuracy
2. Validate legal compliance
3. Approve changes before filing
4. Mark reviews as complete using: `processor.mark_review_complete()`

---

*Generated by Affidavit Enhancement Processor*
"""
        
        return report

    def _requires_review(self, updates: List[EvidenceUpdate]) -> bool:
        """Determine if affidavit requires manual review based on evidence updates."""
        critical_count = sum(1 for u in updates if u.priority == 'critical')
        high_count = sum(1 for u in updates if u.priority == 'high')
        
        review_settings = self.config.get("review_settings", {})
        
        # Require review for any critical updates if enabled
        if review_settings.get("require_review_for_critical", True) and critical_count > 0:
            return True
            
        # Require review for multiple high priority updates if enabled
        high_threshold = review_settings.get("high_priority_threshold", 3)
        if review_settings.get("require_review_for_multiple_high", True) and high_count >= high_threshold:
            return True
            
        return False

    def _create_review_status(self, affidavit_path: Path, enhanced_path: Path, 
                            updates: List[EvidenceUpdate]) -> ReviewStatus:
        """Create a review status entry for an enhanced affidavit."""
        critical_count = sum(1 for u in updates if u.priority == 'critical')
        high_count = sum(1 for u in updates if u.priority == 'high')
        
        return ReviewStatus(
            affidavit_path=str(affidavit_path),
            enhanced_path=str(enhanced_path),
            critical_updates=critical_count,
            high_priority_updates=high_count,
            total_updates=len(updates),
            requires_review=self._requires_review(updates)
        )

    def _save_review_status(self, review_status: ReviewStatus):
        """Save review status to file for tracking."""
        review_file = self.review_dir / f"{Path(review_status.affidavit_path).stem}_review_status.json"
        
        with open(review_file, 'w') as f:
            json.dump(asdict(review_status), f, indent=2)
        
        logger.info(f"📝 Review status saved: {review_file}")

    def _generate_review_alert(self, review_status: ReviewStatus) -> str:
        """Generate an alert message for documents requiring review."""
        if not review_status.requires_review:
            return ""
            
        alert = """
🚨 **LEGAL REVIEW REQUIRED** 🚨

**Document**: {Path(review_status.affidavit_path).name}
**Enhanced Version**: {Path(review_status.enhanced_path).name}

**Evidence Updates Summary:**
- Critical Updates: {review_status.critical_updates}
- High Priority Updates: {review_status.high_priority_updates}
- Total Updates: {review_status.total_updates}

**⚠️ WARNING**: This document contains critical evidence updates and MUST be reviewed by qualified legal counsel before filing or submission.

**Required Actions:**
1. Review enhanced affidavit for accuracy
2. Validate legal compliance
3. Approve changes before filing
4. Mark review as complete

**Status**: {review_status.approval_status.upper()}
**Ready for Filing**: {'✅ YES' if review_status.ready_for_filing else '❌ NO'}
"""
        return alert

    def mark_review_complete(self, affidavit_path: str, reviewer: str, 
                           approval_status: str, comments: str = "") -> bool:
        """Mark an affidavit review as complete."""
        review_file = self.review_dir / f"{Path(affidavit_path).stem}_review_status.json"
        
        if not review_file.exists():
            logger.error(f"Review file not found: {review_file}")
            return False
            
        # Load existing status
        with open(review_file, 'r') as f:
            status_dict = json.load(f)
        
        # Update review status
        status_dict.update({
            'review_completed': True,
            'reviewer': reviewer,
            'review_date': datetime.now().isoformat(),
            'approval_status': approval_status,
            'review_comments': comments,
            'ready_for_filing': approval_status == 'approved'
        })
        
        # Save updated status
        with open(review_file, 'w') as f:
            json.dump(status_dict, f, indent=2)
            
        logger.info(f"✅ Review completed for {Path(affidavit_path).name} by {reviewer}")
        return True

    def get_pending_reviews(self) -> List[ReviewStatus]:
        """Get list of affidavits pending review."""
        pending_reviews = []
        
        for review_file in self.review_dir.glob("*_review_status.json"):
            with open(review_file, 'r') as f:
                status_dict = json.load(f)
            
            status = ReviewStatus(**status_dict)
            if status.requires_review and not status.review_completed:
                pending_reviews.append(status)
                
        return pending_reviews

    def generate_review_summary_report(self) -> str:
        """Generate a comprehensive review summary report."""
        pending_reviews = self.get_pending_reviews()
        
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        total_critical = sum(r.critical_updates for r in pending_reviews)
        total_high = sum(r.high_priority_updates for r in pending_reviews)
        total_updates = sum(r.total_updates for r in pending_reviews)
        
        report = """# 🚨 Critical Affidavit Review Alert

**Timestamp**: {timestamp}
**Trigger**: Critical evidence updates detected
**Status**: LEGAL REVIEW REQUIRED

## Evidence Analysis Summary

- **Critical Updates**: {total_critical}
- **High Priority Updates**: {total_high}  
- **Total Updates**: {total_updates}

## Action Required

Critical evidence updates have been automatically incorporated into affidavits.
**Legal review is strongly recommended** before filing or submitting these documents.

## Files Requiring Review

"""
        
        if not pending_reviews:
            report += "✅ No affidavits currently require review.\n"
        else:
            for i, review in enumerate(pending_reviews, 1):
                report += """### {i}. {Path(review.affidavit_path).name}

**Enhanced File**: `{Path(review.enhanced_path).name}`
**Critical Updates**: {review.critical_updates}
**High Priority Updates**: {review.high_priority_updates}
**Total Updates**: {review.total_updates}
**Status**: {review.approval_status.upper()}

"""

        report += """
## Enhanced Files Location

Enhanced affidavits can be found in the `{self.output_dir}` directory.
Original documents have been backed up to `{self.backup_dir}`.

## Review Status Files

Review tracking files are saved in: `{self.review_dir}`

## Next Steps

1. Review enhanced affidavits for accuracy
2. Validate legal compliance  
3. Approve changes before filing
4. Mark reviews as complete using: `processor.mark_review_complete()`
5. Close this issue once review is complete

---

*This report was automatically generated due to critical evidence updates.*
"""
        
        return report


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Enhance affidavits with evidence updates")
    parser.add_argument("--config", help="Configuration file path")
    parser.add_argument("--since", help="Only process changes since timestamp (ISO format)")
    parser.add_argument("--verbose", "-v", action="store_true", help="Verbose logging")
    
    args = parser.parse_args()
    
    # Set up logging
    level = logging.DEBUG if args.verbose else logging.INFO
    logging.basicConfig(level=level, format='%(asctime)s - %(levelname)s - %(message)s')
    
    # Process affidavits
    processor = AffidavitProcessor(args.config)
    results = processor.process_all_affidavits(args.since)
    
    # Generate and save report
    report = processor.generate_enhancement_report(results)
    
    report_path = Path("AFFIDAVIT_ENHANCEMENT_REPORT.md")
    report_path.write_text(report)
    
    print(f"Enhancement complete. Report saved to: {report_path}")
    print(f"Enhanced files in: {processor.output_dir}")